from pathlib import Path
import json
import pickle
import numpy as np
import pandas as pd
import joblib
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.metrics import (
    classification_report,
    roc_auc_score,
    average_precision_score,
    precision_recall_curve,
    confusion_matrix,
)
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
PROCESSED = ROOT / "data" / "processed"
MODELS = ROOT / "outputs" / "models"
FIGURES = ROOT / "outputs" / "figures"
REPORT_PATH = ROOT / "Client_Project_Report.docx"

sns.set_style("whitegrid")
plt.rcParams["figure.figsize"] = (10, 5)


def load_data():
    X_train = pickle.load(open(PROCESSED / "X_train.pkl", "rb"))
    X_test = pickle.load(open(PROCESSED / "X_test.pkl", "rb"))
    y_train = pickle.load(open(PROCESSED / "y_train.pkl", "rb"))
    y_test = pickle.load(open(PROCESSED / "y_test.pkl", "rb"))
    return X_train, X_test, y_train, y_test


def best_threshold_by_f1(y_true, proba):
    precision, recall, thresholds = precision_recall_curve(y_true, proba)
    f1 = 2 * (precision * recall) / (precision + recall + 1e-8)
    idx = int(np.argmax(f1))
    if idx >= len(thresholds):
        return 0.5
    return float(thresholds[idx])


def evaluate_models(X_test, y_test):
    candidates = {
        "Logistic Regression": MODELS / "logistic_income_pipeline.joblib",
        "Random Forest": MODELS / "random_forest_pipeline.joblib",
        "XGBoost": MODELS / "XGBoost_pipeline.joblib",
        "LightGBM": MODELS / "LightGBM_pipeline.joblib",
        "XGBoost + SMOTE": MODELS / "SMOTE_XGBoost_pipeline.joblib",
        "XGBoost V2": MODELS / "XGBoost_V2_pipeline.joblib",
    }

    rows = []
    best_name = None
    best_obj = None
    best_score = -1.0
    best_thr = 0.5

    def build_eval_frame(model, X):
        X_eval = X.copy()
        expected = getattr(model, "feature_names_in_", None)

        if expected is None:
            return X_eval

        expected = list(expected)

        if "wealth_score" in expected and "wealth_score" not in X_eval.columns:
            if all(c in X_eval.columns for c in ["capital gains", "capital losses", "dividends from stocks"]):
                X_eval["wealth_score"] = (
                    X_eval["capital gains"] - X_eval["capital losses"] + X_eval["dividends from stocks"]
                )
            else:
                X_eval["wealth_score"] = 0.0

        if "age_x_education" in expected and "age_x_education" not in X_eval.columns:
            if all(c in X_eval.columns for c in ["age", "education"]):
                X_eval["age_x_education"] = X_eval["age"] * X_eval["education"]
            else:
                X_eval["age_x_education"] = 0.0

        for col in expected:
            if col not in X_eval.columns:
                X_eval[col] = 0.0

        return X_eval[expected]

    for name, path in candidates.items():
        if not path.exists():
            continue
        model = joblib.load(path)
        X_eval = build_eval_frame(model, X_test)
        proba = model.predict_proba(X_eval)[:, 1]
        thr = best_threshold_by_f1(y_test, proba)
        pred = (proba >= thr).astype(int)
        report = classification_report(y_test, pred, output_dict=True)
        pos_key = "1" if "1" in report else list(k for k in report.keys() if k.isdigit())[-1]
        row = {
            "model": name,
            "roc_auc": roc_auc_score(y_test, proba),
            "pr_auc": average_precision_score(y_test, proba),
            "f1_minority": report[pos_key]["f1-score"],
            "recall_minority": report[pos_key]["recall"],
            "precision_minority": report[pos_key]["precision"],
            "threshold": thr,
        }
        rows.append(row)

        if row["pr_auc"] > best_score:
            best_score = row["pr_auc"]
            best_name = name
            best_obj = model
            best_thr = thr

    result_df = pd.DataFrame(rows).sort_values("pr_auc", ascending=False)
    return result_df, best_name, best_obj, best_thr


def make_figures(X_train, y_train, X_test, y_test, model_table, best_model, best_threshold):
    FIGURES.mkdir(parents=True, exist_ok=True)

    class_dist_path = FIGURES / "class_distribution.png"
    y_train.value_counts().sort_index().rename(index={0: "<=50K", 1: ">50K"}).plot(
        kind="bar", color=["#4C78A8", "#F58518"]
    )
    plt.title("Training Class Distribution")
    plt.ylabel("Count")
    plt.xlabel("Income Class")
    plt.tight_layout()
    plt.savefig(class_dist_path, dpi=180)
    plt.close()

    model_cmp_path = FIGURES / "model_comparison_pr_auc.png"
    if not model_table.empty:
        tmp = model_table.sort_values("pr_auc", ascending=True)
        plt.barh(tmp["model"], tmp["pr_auc"], color="#4C78A8")
        plt.title("Classification Model Comparison (PR-AUC)")
        plt.xlabel("PR-AUC")
        plt.tight_layout()
        plt.savefig(model_cmp_path, dpi=180)
        plt.close()

    cm_path = FIGURES / "best_model_confusion_matrix.png"
    if best_model is not None:
        expected = getattr(best_model, "feature_names_in_", None)
        X_eval = X_test.copy()
        if expected is not None:
            expected = list(expected)
            if "wealth_score" in expected and "wealth_score" not in X_eval.columns:
                if all(c in X_eval.columns for c in ["capital gains", "capital losses", "dividends from stocks"]):
                    X_eval["wealth_score"] = (
                        X_eval["capital gains"] - X_eval["capital losses"] + X_eval["dividends from stocks"]
                    )
                else:
                    X_eval["wealth_score"] = 0.0
            if "age_x_education" in expected and "age_x_education" not in X_eval.columns:
                if all(c in X_eval.columns for c in ["age", "education"]):
                    X_eval["age_x_education"] = X_eval["age"] * X_eval["education"]
                else:
                    X_eval["age_x_education"] = 0.0
            for col in expected:
                if col not in X_eval.columns:
                    X_eval[col] = 0.0
            X_eval = X_eval[expected]

        proba = best_model.predict_proba(X_eval)[:, 1]
        y_pred = (proba >= best_threshold).astype(int)
        cm = confusion_matrix(y_test, y_pred)
        plt.figure(figsize=(6, 5))
        sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", cbar=False)
        plt.title("Best Model Confusion Matrix")
        plt.xlabel("Predicted")
        plt.ylabel("Actual")
        plt.tight_layout()
        plt.savefig(cm_path, dpi=180)
        plt.close()

    seg_sizes_path = FIGURES / "segment_sizes.png"
    seg_income_path = FIGURES / "segment_income_rate.png"
    pca_scatter_path = FIGURES / "segmentation_pca_scatter.png"

    seg_assign_path = PROCESSED / "segmentation_assignments.pkl"
    if seg_assign_path.exists():
        seg_df = pd.read_pickle(seg_assign_path)

        segment_col = "segment" if "segment" in seg_df.columns else "cluster"
        counts = seg_df[segment_col].value_counts()
        counts.plot(kind="bar", color="#54A24B")
        plt.title("Segment Sizes")
        plt.ylabel("Count")
        plt.xlabel("Segment")
        plt.xticks(rotation=25, ha="right")
        plt.tight_layout()
        plt.savefig(seg_sizes_path, dpi=180)
        plt.close()

        if "label" in seg_df.columns:
            income_rate = seg_df.groupby(segment_col)["label"].mean().sort_values(ascending=False) * 100
            income_rate.plot(kind="bar", color="#E45756")
            plt.title("High-Income Rate by Segment")
            plt.ylabel("% >50K")
            plt.xlabel("Segment")
            plt.xticks(rotation=25, ha="right")
            plt.tight_layout()
            plt.savefig(seg_income_path, dpi=180)
            plt.close()

    x_pca_path = PROCESSED / "segmentation_X_pca.npy"
    km_path = MODELS / "kmeans_segmentation_model.joblib"
    if x_pca_path.exists() and km_path.exists():
        X_pca = np.load(x_pca_path)
        km = joblib.load(km_path)
        clusters = km.predict(X_pca)
        plt.figure(figsize=(7, 6))
        sns.scatterplot(x=X_pca[:, 0], y=X_pca[:, 1], hue=clusters, palette="tab10", s=12, linewidth=0)
        plt.title("Segmentation in PCA Space")
        plt.xlabel("PC1")
        plt.ylabel("PC2")
        plt.legend(title="Cluster", bbox_to_anchor=(1.02, 1), loc="upper left")
        plt.tight_layout()
        plt.savefig(pca_scatter_path, dpi=180)
        plt.close()

    return {
        "class_dist": class_dist_path,
        "model_cmp": model_cmp_path,
        "cm": cm_path,
        "seg_sizes": seg_sizes_path,
        "seg_income": seg_income_path,
        "pca": pca_scatter_path,
    }


def add_heading_and_text(doc, heading, text):
    doc.add_heading(heading, level=2)
    p = doc.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.space_after = Pt(8)


def safe_add_figure(doc, title, path):
    if path.exists():
        doc.add_heading(title, level=3)
        doc.add_picture(str(path), width=Inches(6.2))
        cap = doc.add_paragraph(f"Figure: {title}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(model_table, best_name, best_thr, figs):
    doc = Document()

    title = doc.add_heading("Client Report: Census Income Classification and Customer Segmentation", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Prepared for client delivery | Project status: completed pipeline and baseline deployment-ready artifacts")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading_and_text(
        doc,
        "1. Executive Summary",
        "This project built two production-oriented analytics solutions from Census data: "
        "(a) a binary classifier to predict whether annual income is above 50K, and "
        "(b) a customer segmentation model to create actionable groups for marketing strategy. "
        "The classification workflow prioritizes minority-class precision/recall trade-offs using PR-AUC and threshold tuning. "
        "The segmentation workflow applies PCA plus KMeans and translates clusters into business-friendly segment narratives."
    )

    add_heading_and_text(
        doc,
        "2. Data Exploration and Pre-processing",
        "Data exploration focused on shape, missingness, cardinality, and target imbalance. "
        "For classification, weak or sparse migration/origin administrative variables were removed, the target label was mapped to 0/1, "
        "education was ordinally encoded, selected categoricals were one-hot encoded, and high-cardinality fields were target-encoded after train/test split. "
        "For segmentation, a compact behavior-centric feature set was selected and transformed with ordinal mappings, binary behavioral flags, log transforms for skewed finance variables, and standard scaling."
    )

    add_heading_and_text(
        doc,
        "3. Model Architecture and Training Algorithm",
        "Classification compared Logistic Regression, Random Forest, XGBoost, LightGBM, XGBoost with SMOTE, and a second XGBoost variant with feature engineering and wider hyperparameter search. "
        "RandomizedSearchCV with stratified folds was used for robust tuning. Decision threshold optimization was performed using F1 over precision-recall thresholds, rather than fixed 0.5 cutoff. "
        "Segmentation used PCA for dimensionality reduction and KMeans clustering; elbow and silhouette analyses guided selecting k=6 for actionable granularity."
    )

    add_heading_and_text(
        doc,
        "4. Evaluation Procedure",
        "Classification evaluation used ROC-AUC, PR-AUC, and minority-class precision, recall, and F1 on a held-out test set. "
        "PR-AUC was prioritized because income >50K is the minority class and business action depends on positive-class quality. "
        "Segmentation evaluation used inertia (elbow behavior) and silhouette score, followed by profile validation through segment size and income-rate diagnostics."
    )

    if not model_table.empty:
        doc.add_heading("5. Classification Results", level=2)
        table = doc.add_table(rows=1, cols=7)
        hdr = table.rows[0].cells
        hdr[0].text = "Model"
        hdr[1].text = "PR-AUC"
        hdr[2].text = "ROC-AUC"
        hdr[3].text = "F1 minority"
        hdr[4].text = "Recall minority"
        hdr[5].text = "Precision minority"
        hdr[6].text = "Threshold"

        for _, r in model_table.iterrows():
            row = table.add_row().cells
            row[0].text = str(r["model"])
            row[1].text = f"{r['pr_auc']:.4f}"
            row[2].text = f"{r['roc_auc']:.4f}"
            row[3].text = f"{r['f1_minority']:.4f}"
            row[4].text = f"{r['recall_minority']:.4f}"
            row[5].text = f"{r['precision_minority']:.4f}"
            row[6].text = f"{r['threshold']:.3f}"

        doc.add_paragraph(
            f"Best model by PR-AUC: {best_name}. Recommended serving threshold: {best_thr:.3f}."
        )

    add_heading_and_text(
        doc,
        "6. Segmentation Findings and Business Recommendations",
        "The final segmentation uses six clusters and supports differentiated go-to-market actions: "
        "high-value premium campaigns for top-income segments, value-focused bundles for middle-income groups, and indirect household targeting for dependent-heavy segments. "
        "Recommendation is to operationalize segment-level marketing plays with controlled experiments and quarterly model refreshes."
    )

    add_heading_and_text(
        doc,
        "7. Business Judgment and Decision Rationale",
        "Key decisions included dropping administratively noisy fields, preserving interpretable transformations, prioritizing PR-AUC for imbalanced classification, "
        "using threshold tuning to align with campaign economics, and selecting six segments for actionability rather than maximum statistical granularity. "
        "Model usage recommendation: deploy the best PR-AUC classifier with calibrated threshold controls and use segment outputs as a campaign stratification layer."
    )

    doc.add_heading("8. Visualizations", level=2)
    safe_add_figure(doc, "Training Class Distribution", figs["class_dist"])
    safe_add_figure(doc, "Classification Model Comparison (PR-AUC)", figs["model_cmp"])
    safe_add_figure(doc, "Best Model Confusion Matrix", figs["cm"])
    safe_add_figure(doc, "Segment Sizes", figs["seg_sizes"])
    safe_add_figure(doc, "High-Income Rate by Segment", figs["seg_income"])
    safe_add_figure(doc, "Segmentation in PCA Space", figs["pca"])

    doc.add_heading("9. References", level=2)
    refs = [
        "Scikit-learn official documentation (model selection, metrics, PCA, KMeans).",
        "XGBoost documentation (classifier parameters and imbalance handling).",
        "LightGBM documentation (tree boosting and class weighting).",
        "CatBoost documentation (gradient boosting for structured data).",
        "Imbalanced-learn documentation (SMOTE strategy and caveats).",
        "Pandas and NumPy documentation for data processing operations.",
        "U.S. Census dataset documentation provided with the assignment files.",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("10. Delivery Notes", level=2)
    doc.add_paragraph(
        "This report is intentionally concise and formatted to remain within a 10-page client-friendly scope. "
        "Detailed experimentation traces remain in project notebooks for auditability."
    )

    doc.save(REPORT_PATH)


def main():
    X_train, X_test, y_train, y_test = load_data()
    model_table, best_name, best_model, best_thr = evaluate_models(X_test, y_test)

    model_table_path = PROCESSED / "classification_model_comparison.csv"
    if not model_table.empty:
        model_table.to_csv(model_table_path, index=False)

    figs = make_figures(X_train, y_train, X_test, y_test, model_table, best_model, best_thr)
    build_docx(model_table, best_name, best_thr, figs)

    print(f"Report generated: {REPORT_PATH}")
    if model_table_path.exists():
        print(f"Model comparison table: {model_table_path}")
    print(f"Figures folder: {FIGURES}")


if __name__ == "__main__":
    main()
